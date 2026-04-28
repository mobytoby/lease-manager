#!/usr/bin/env python3
"""Generate a complete lease (core + addenda) from JSON profiles.

Usage:
    python3 generate_lease.py <landlord.json> <property.json> <tenancy.json> <output.docx>

This script is the assembly engine for the lease wizard. The wizard's job will be
to collect inputs, write the JSON files, and call this generator. The same generator
handles new leases, renewals, with/without pets, fixed/month-to-month, etc. — driven
entirely by the JSON inputs.

Section-by-section conditional logic lives here. Variable substitution is a plain
Python format-string operation; the JSON profiles are the single source of truth.
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from lease_helpers import (
    fmt_date, fmt_money, fmt_address, join_names,
    add_heading, add_para, add_signature_block, set_cell_shading,
)


# ---------------------------------------------------------------------------
# Numbering helpers
# ---------------------------------------------------------------------------

def _get_style_num_id(doc, style_name="List Number"):
    """Return the numId string baked into the named paragraph style."""
    for style in doc.styles:
        if style.name == style_name:
            pPr = style.element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    numId_elem = numPr.find(qn("w:numId"))
                    if numId_elem is not None:
                        return numId_elem.get(qn("w:val"))
    return None


def _alloc_num_id(doc):
    """Create a fresh w:num entry (startOverride=1) in the numbering part and
    return its numId string.  All paragraphs in a single list must share this
    id so Word treats them as one counting sequence.
    """
    style_num_id = _get_style_num_id(doc)
    if style_num_id is None:
        return None

    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return None
    numbering_elem = numbering_part._element

    # Resolve the abstractNumId behind the style's numId
    abstract_num_id = None
    for num in numbering_elem.findall(qn("w:num")):
        if num.get(qn("w:numId")) == style_num_id:
            ref = num.find(qn("w:abstractNumId"))
            if ref is not None:
                abstract_num_id = ref.get(qn("w:val"))
            break
    if abstract_num_id is None:
        return None

    # Allocate a new numId above the current maximum
    existing_ids = [
        int(n.get(qn("w:numId")))
        for n in numbering_elem.findall(qn("w:num"))
        if n.get(qn("w:numId")) is not None
    ]
    new_id = str(max(existing_ids) + 1)

    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), new_id)
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), abstract_num_id)
    new_num.append(abs_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    new_num.append(lvl_override)
    numbering_elem.append(new_num)
    return new_id


def _apply_num_id(paragraph, num_id, ilvl="0"):
    """Write an explicit paragraph-level numPr so *paragraph* uses *num_id*.

    python-docx stores numId in the style definition rather than in each
    paragraph's XML, so we must inject a paragraph-level override to redirect
    a specific paragraph to a different numId.
    """
    if num_id is None:
        return
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
    # ilvl must come before numId in the schema
    ilvl_elem = numPr.find(qn("w:ilvl"))
    if ilvl_elem is None:
        ilvl_elem = OxmlElement("w:ilvl")
        numPr.insert(0, ilvl_elem)
    ilvl_elem.set(qn("w:val"), ilvl)
    numId_elem = numPr.find(qn("w:numId"))
    if numId_elem is None:
        numId_elem = OxmlElement("w:numId")
        numPr.append(numId_elem)
    numId_elem.set(qn("w:val"), num_id)


def numbered_list(doc, items, font_size=10):
    """Add *items* as a fresh numbered list that always starts at 1."""
    num_id = _alloc_num_id(doc)
    paragraphs = []
    for text in items:
        p = doc.add_paragraph(text, style="List Number")
        p.runs[0].font.size = Pt(font_size)
        _apply_num_id(p, num_id)
        paragraphs.append(p)
    return paragraphs


# ---------------------------------------------------------------------------
# Top-level
# ---------------------------------------------------------------------------

def generate(landlord, property_, tenancy, output_path):
    ctx = build_context(landlord, property_, tenancy)
    doc = Document()
    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    write_cover(doc, ctx)
    write_table_of_contents(doc, ctx)
    write_renewal_recital_if_applicable(doc, ctx)
    write_lease_preamble(doc, ctx)
    write_section_1_basic_terms(doc, ctx)
    write_section_2_additional_terms(doc, ctx)
    write_contact_info(doc, ctx)
    write_main_signature(doc, ctx)

    # Addenda — order mirrors the Table of Contents
    if ctx["addenda"]["pets"]["include"]:
        write_pet_addendum(doc, ctx)
    if ctx["addenda"]["parking"]["include"]:
        write_parking_addendum(doc, ctx)
    if ctx["addenda"]["rules"]["include"]:
        write_rules_addendum(doc, ctx)
    if ctx["addenda"]["lawn_care"]["include"]:
        write_lawn_care_addendum(doc, ctx)
    if ctx["addenda"]["lead_paint_disclosure"]["include"]:
        write_lead_paint_disclosure(doc, ctx)
    if ctx["addenda"]["mold_disclosure"]["include"]:
        write_mold_addendum(doc, ctx)
    if ctx["addenda"]["fire_safety"]["include"]:
        write_fire_safety_addendum(doc, ctx)

    add_initials_footer(doc, ctx["tenancy"]["tenants"])
    doc.save(output_path)


def add_initials_footer(doc, tenants):
    """Add a right-aligned tenant initials line to the footer of every page."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if len(tenants) >= 2:
        first1 = tenants[0]["name"].split()[0]
        first2 = tenants[1]["name"].split()[0]
        text = f"{first1} Initials: _______     {first2} Initials: _______"
    else:
        first = tenants[0]["name"].split()[0] if tenants else "Tenant"
        text = f"{first} Initials: _______"

    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = None  # inherit theme colour (dark grey in most themes)


# ---------------------------------------------------------------------------
# Context building — flatten the three JSON profiles into one easy-to-use dict
# ---------------------------------------------------------------------------

def build_context(landlord, property_, tenancy):
    tenant_names = [t["name"] for t in tenancy["tenants"]]
    addr = property_["address"]
    util = property_["utilities"]

    # Telephone/cable/internet setup labels
    def utility_setup_label(u, name):
        s = u.get("setup", "tenant")
        if s == "tenant":
            return f"Tenant will arrange and pay for the cost of {name}."
        if s == "landlord_pay":
            return f"Landlord will arrange and pay for the cost of {name}."
        if s == "landlord_fee":
            fee = u.get("fee", 0)
            return (
                f"Landlord will arrange {name} and Tenant will pay a fixed monthly "
                f"fee of {fmt_money(fee)} as Additional Rent."
            )
        return f"Tenant will arrange and pay for the cost of {name}."

    return {
        "landlord": landlord,
        "property": property_,
        "tenancy": tenancy,
        "tenant_names": tenant_names,
        "tenant_names_joined": join_names(tenant_names),
        "property_address_oneline": fmt_address(addr),
        "utility_label": utility_setup_label,
        "addenda": tenancy["addenda"],
        "is_renewal": tenancy.get("lease_kind") == "renewal",
        "term_is_fixed": tenancy["term"]["type"] == "fixed",
    }


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def write_cover(doc, ctx):
    add_para(doc, "", size=10)
    add_para(doc, "", size=10)
    add_para(doc, "", size=10)
    add_heading(doc, "RESIDENTIAL LEASE AGREEMENT", level=0, center=True)
    add_para(doc, "", size=10)
    add_para(doc, ctx["property_address_oneline"], size=14, align="center", bold=True)
    add_para(doc, "", size=10)
    add_para(doc, "", size=10)
    add_para(
        doc,
        f"Landlord: {ctx['landlord']['name']}",
        size=12, align="center",
    )
    add_para(
        doc,
        f"Tenant(s): {ctx['tenant_names_joined']}",
        size=12, align="center",
    )
    if ctx["term_is_fixed"]:
        add_para(
            doc,
            f"Term: {fmt_date(ctx['tenancy']['term']['start_date'])} "
            f"through {fmt_date(ctx['tenancy']['term']['end_date'])}",
            size=12, align="center",
        )
    else:
        add_para(
            doc,
            f"Term: month-to-month beginning {fmt_date(ctx['tenancy']['term']['start_date'])}",
            size=12, align="center",
        )
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Table of contents
# ---------------------------------------------------------------------------

def write_table_of_contents(doc, ctx):
    add_heading(doc, "Table of Contents", level=1)
    add_para(doc, "The following documents are incorporated into and attached as part of this Lease:")
    add_para(doc, "", size=4)
    items = ["Basic Terms & Additional Terms"]
    a = ctx["addenda"]
    if a["pets"]["include"]:
        items.append("Pet Addendum")
    if a["parking"]["include"]:
        items.append("Parking Addendum")
    if a["rules"]["include"]:
        items.append("Rules Addendum")
    if a["lawn_care"]["include"]:
        items.append("Tenant Yard Care Addendum")
    if a["lead_paint_disclosure"]["include"]:
        items.append("Lead-Based Paint Hazard Disclosure (with EPA pamphlet)")
    if a["mold_disclosure"]["include"]:
        items.append("Mold Addendum (with mold pamphlet)")
    if a["fire_safety"]["include"]:
        items.append("Fire Safety Addendum")
    for item in items:
        p = doc.add_paragraph(item, style=None)
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Renewal recital
# ---------------------------------------------------------------------------

def write_renewal_recital_if_applicable(doc, ctx):
    if not ctx["is_renewal"]:
        return
    add_heading(doc, "Renewal Recital", level=1)
    prior = ctx["tenancy"].get("prior_lease_effective_date")
    add_para(
        doc,
        f"This Lease is a renewal of the prior residential lease for the Property "
        f"between Landlord and Tenant dated {fmt_date(prior)} (the Prior Lease). "
        f"This Lease supersedes and replaces the Prior Lease as of the Start Date set "
        f"forth below. Except as expressly provided herein, the parties intend for "
        f"the tenancy to continue uninterrupted, and Tenant's possession of the "
        f"Property is continuous from the Prior Lease.",
    )
    sd = ctx["tenancy"]["security_deposit"]
    if sd.get("already_held"):
        add_para(
            doc,
            f"The Security Deposit of {fmt_money(sd['amount'])} previously paid by "
            f"Tenant under the Prior Lease shall continue to be held by Landlord "
            f"under the terms of this Lease, and no additional deposit is required "
            f"at signing of this Lease.",
        )
    add_para(doc, "", size=4)


# ---------------------------------------------------------------------------
# Preamble
# ---------------------------------------------------------------------------

def write_lease_preamble(doc, ctx):
    add_para(
        doc,
        f"This Residential Lease (Lease) is entered into on the date of the last "
        f"signature below (the Effective Date) between {ctx['landlord']['name']} "
        f"(Landlord) and {ctx['tenant_names_joined']} (together and severally, "
        f"Tenant) for the property located at {ctx['property_address_oneline']} "
        f"(Property).",
    )
    add_para(
        doc,
        "Landlord hereby leases the Property to Tenant, subject to the terms and "
        "conditions of this Lease.",
    )


# ---------------------------------------------------------------------------
# Section 1: Basic Terms
# ---------------------------------------------------------------------------

def write_section_1_basic_terms(doc, ctx):
    add_heading(doc, "1. Basic Terms", level=1)

    write_1_1_amounts_due_upfront(doc, ctx)
    write_1_2_occupants(doc, ctx)
    write_1_3_property(doc, ctx)
    write_1_4_term(doc, ctx)
    write_1_5_rent(doc, ctx)
    write_1_6_utilities(doc, ctx)
    write_1_7_security_deposit(doc, ctx)
    write_1_8_tenant_insurance(doc, ctx)
    write_1_9_parking(doc, ctx)
    write_1_10_storage(doc, ctx)
    write_1_11_pets(doc, ctx)
    write_1_12_additional_rules(doc, ctx)
    write_1_13_smoking(doc, ctx)


def write_1_1_amounts_due_upfront(doc, ctx):
    add_heading(doc, "1.1. Amounts Due From Tenant Upfront", level=2)

    sd = ctx["tenancy"]["security_deposit"]
    rent = ctx["tenancy"]["rent"]
    schedule = sd.get("payment_schedule", "lump_sum")

    add_para(doc, "1.1.1. Refundable Deposit", bold=True)
    if sd.get("already_held"):
        add_para(
            doc,
            f"{fmt_money(sd['amount'])} Security Deposit (held by Landlord from prior "
            f"lease; no additional payment due at signing of this Lease).",
        )
    elif schedule == "installments":
        inst = sd.get("installments", {})
        first_due_at = inst.get("first_due", "start")
        when_label = "Start Date" if first_due_at == "start" else "signing"
        add_para(
            doc,
            f"{fmt_money(sd['amount'])} Security Deposit, payable in "
            f"{inst.get('count', 0)} consecutive equal monthly installments of "
            f"{fmt_money(inst.get('amount', 0))} beginning at {when_label} "
            f"(see Section 1.7).",
        )
    else:
        add_para(
            doc,
            f"{fmt_money(sd['amount'])} Security Deposit due at signing "
            f"(see Section 1.7).",
        )

    add_para(doc, "1.1.2. Rent for First Month", bold=True)
    add_para(
        doc,
        f"{fmt_money(rent.get('prorated_first_month', 0))} Prorated Monthly Rent for "
        f"partial first month of the Term due at Start Date (see Section 1.5).",
    )
    add_para(
        doc,
        f"{fmt_money(rent['monthly_rent'])} Monthly Rent for full first month of the "
        f"Term due on the first Monthly Rent Due Date after the Start Date "
        f"(see Section 1.5).",
    )

    add_para(doc, "1.1.3. Total Due Upfront", bold=True)
    due_at_signing = sd.get("additional_due_at_signing", 0)
    due_at_start = rent.get("prorated_first_month", 0)
    if not sd.get("already_held"):
        if schedule == "installments":
            inst = sd.get("installments", {})
            first_due_at = inst.get("first_due", "start")
            first_amount = inst.get("amount", 0)
            if first_due_at == "signing":
                due_at_signing += first_amount
            else:
                due_at_start += first_amount
        else:
            due_at_signing += sd["amount"]
    add_para(doc, f"{fmt_money(due_at_signing)} due at signing.")
    add_para(doc, f"{fmt_money(due_at_start)} due at Start Date.")


def write_1_2_occupants(doc, ctx):
    add_heading(doc, "1.2. Additional Occupant Information", level=2)
    occ = ctx["tenancy"].get("additional_occupants", [])
    add_para(doc, "The following additional occupants (Occupants) may occupy the Property:")
    if not occ:
        add_para(doc, "None.")
        return
    table = doc.add_table(rows=1 + len(occ), cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].paragraphs[0].add_run("Name").bold = True
    hdr[1].paragraphs[0].add_run("Age").bold = True
    for i, o in enumerate(occ, start=1):
        cells = table.rows[i].cells
        cells[0].text = o.get("name", "")
        cells[1].text = "" if o.get("age") in (None, "") else str(o["age"])


def write_1_3_property(doc, ctx):
    add_heading(doc, "1.3. Property", level=2)
    p = ctx["property"]
    add_para(
        doc,
        f"The Property is a {p['property_type']} located at "
        f"{ctx['property_address_oneline']}.",
    )
    appliances = [a["name"] for a in p.get("appliances", [])]
    if appliances:
        add_para(
            doc,
            "The Property includes all appliances, fixtures, and equipment installed "
            "on the Property as of the Start Date, including the following: "
            + "; ".join(appliances) + ".",
        )


def write_1_4_term(doc, ctx):
    add_heading(doc, "1.4. Term", level=2)
    t = ctx["tenancy"]["term"]
    if ctx["term_is_fixed"]:
        add_para(
            doc,
            f"Fixed Term. The term of this Lease (Term) will begin on "
            f"{fmt_date(t['start_date'])} (Start Date) and end on "
            f"{fmt_date(t['end_date'])} (Expiration Date).",
        )
        notice = ctx["landlord"]["defaults"]["month_to_month_tenant_notice_days"]
        add_para(
            doc,
            f"If Landlord accepts Tenant's payment of the Monthly Rent for the month "
            f"after the end of the Term, this Lease will continue on a month-to-month "
            f"basis on the same terms and conditions. Tenant may terminate the "
            f"month-to-month tenancy by providing at least {notice} days' written "
            f"notice to Landlord, with Rent prorated based on a 30-day month. "
            f"Landlord may terminate the month-to-month tenancy only for just cause "
            f"and must provide advance written notice in accordance with applicable "
            f"law.",
        )
    else:
        notice = ctx["landlord"]["defaults"]["month_to_month_tenant_notice_days"]
        add_para(
            doc,
            f"Month-to-Month Term. The term of this Lease (Term) will begin on "
            f"{fmt_date(t['start_date'])} (Start Date) and continue on a month-to-"
            f"month basis until terminated. Tenant may terminate by providing at "
            f"least {notice} days' written notice to Landlord, with Rent prorated "
            f"based on a 30-day month. Landlord may terminate only for just cause "
            f"and must provide advance written notice in accordance with applicable "
            f"law.",
        )

    # WA RCW 59.18.575 — Early termination for victims of domestic violence,
    # sexual assault, unlawful harassment, or stalking. Required disclosure.
    add_para(
        doc,
        "If Tenant is a victim of family violence or reasonably believes it is "
        "necessary to vacate the Property due to fear of imminent domestic violence; "
        "if Tenant is a victim of sexual assault, or is the parent or guardian of a "
        "dependent who is a victim of sexual assault; or if Tenant is a victim of "
        "unlawful harassment or stalking and reasonably believes it is necessary to "
        "vacate the Property due to fear of imminent harm of such sexual assault, "
        "domestic violence, or unlawful harassment, Tenant may terminate this Lease "
        "by providing Landlord thirty (30) days' prior written notice. If Tenant "
        "provides written notice of termination for any of these reasons, and the "
        "notice includes a valid order for protection or a written record of report "
        "to a qualified third party (such as the police), then Tenant will be "
        "automatically discharged from the payment of Rent following the month in "
        "which the notice was given, and Tenant will be entitled to the return of "
        "the Security Deposit (if applicable), minus any deductions by Landlord "
        "pursuant to Section 2.4 of this Lease.",
    )


def write_1_5_rent(doc, ctx):
    add_heading(doc, "1.5. Rent", level=2)
    rent = ctx["tenancy"]["rent"]
    due_day = rent.get("monthly_rent_due_day", 1)
    grace = rent.get("grace_period_days", 5)
    late_fee = rent.get("late_fee", 0)

    add_para(doc, "1.5.1. Base Rent; Monthly Rent", bold=True)
    add_para(
        doc,
        f"Tenant is responsible for paying monthly rent for the use and occupancy "
        f"of the Property (Base Rent) and all other fixed rent and fixed charges "
        f"described in this Lease (collectively, Monthly Rent) on the {ordinal(due_day)} "
        f"day of each month (Monthly Rent Due Date). The Monthly Rent is "
        f"{fmt_money(rent['monthly_rent'])}. The first Monthly Rent payment is due "
        f"on the first Monthly Rent Due Date after the Start Date.",
    )

    add_para(doc, "1.5.2. Additional Rent", bold=True)
    window = ctx["landlord"]["defaults"]["variable_charge_payment_window_days"]
    add_para(
        doc,
        f"Any amount Tenant may be required to pay Landlord or any other party under "
        f"this Lease in addition to Monthly Rent will be additional rent (Additional "
        f"Rent). The Monthly Rent and any Additional Rent are collectively referred "
        f"to as Rent. Additional Rent includes, without limitation:",
    )
    bullets = [
        f"Any applicable charges for utilities and/or other services to the Property "
        f"that vary by month (Variable Charges), payable to Landlord within {window} "
        f"days of billing.",
        "The cost of utilities required to be arranged for and paid by Tenant directly to the service provider.",
        "Insufficient Funds Fee: variable.",
        f"Late Fee: {fmt_money(late_fee)}.",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    add_para(doc, "1.5.3. Manner of Payment", bold=True)
    add_para(doc, "All Rent payable to Landlord must be paid by one of the following:")

    pm = ctx["landlord"]["payment_methods"]
    override = ctx["tenancy"].get("payment_methods_override")
    if override is not None:
        # Tenancy explicitly restricts/overrides the landlord defaults.
        check_enabled = bool(override.get("check_or_money_order"))
        electronic_methods = override.get("electronic", [])
    else:
        check_enabled = pm["check_or_money_order"]["enabled"]
        electronic_methods = pm.get("electronic", [])

    if check_enabled:
        co = pm["check_or_money_order"]
        add_para(
            doc,
            f"• Check or money order made payable to {co['payable_to']} and delivered "
            f"or mailed to: {fmt_address(co['mail_to'])}.",
        )
    for em in electronic_methods:
        line = f"• Electronically by online payment service: {em['service']}."
        if em.get("notes"):
            line = line + " " + em["notes"]
        add_para(doc, line)


def write_1_6_utilities(doc, ctx):
    add_heading(doc, "1.6. Utilities and Services", level=2)
    add_para(
        doc,
        "Neither Landlord nor Tenant may intentionally cause termination of any "
        "utility services to the Property, including water, heat, electricity, or "
        "gas, except for an interruption of utility services for a reasonable time "
        "in order to make necessary repairs or as a result of normal occupancy of "
        "the Property. Landlord and Tenant agree that utilities and other services "
        "will be provided and paid for as outlined below.",
    )
    util = ctx["property"]["utilities"]
    label = ctx["utility_label"]

    sub = []

    # Electricity
    e = util.get("electricity", {})
    if e.get("paid_by") == "tenant":
        sub.append(("Electricity", f"Tenant will arrange and pay for electrical service "
                                    f"directly to the service provider ({e.get('provider', 'utility')})."))
    else:
        sub.append(("Electricity", "Landlord will provide electrical service; cost is included in Base Rent."))

    # Natural Gas
    g = util.get("natural_gas", {})
    if g.get("paid_by") == "tenant":
        sub.append(("Natural Gas", f"Tenant will arrange and pay for natural gas service "
                                    f"directly to the service provider ({g.get('provider', 'utility')})."))
    elif g.get("paid_by") == "landlord":
        if g.get("included_in_rent"):
            sub.append(("Natural Gas", "Landlord will provide natural gas service; cost is included in Base Rent."))
        else:
            sub.append(("Natural Gas", "Landlord will provide natural gas service; Tenant will reimburse Landlord as Variable Charges."))
    else:
        sub.append(("Natural Gas", "Not applicable."))

    # Heat
    h = util.get("heating", {})
    if h:
        sub.append((
            "Heat",
            f"Heat is provided by {h.get('type', 'the heating system')} and is included "
            f"in the cost of the applicable utility above. The thermostat is a "
            f"{h.get('thermostat', {}).get('brand', '')} unit."
            if h.get("thermostat") else
            "Heat is included in the cost of the applicable utility above.",
        ))

    # Water/Sewer
    w = util.get("water_sewer", {})
    if w.get("included_in_rent"):
        sub.append(("Water and Sewer", "Landlord will provide water and sewer service; cost is included in Base Rent."))
    elif w.get("paid_by") == "tenant":
        sub.append(("Water and Sewer", f"Tenant will arrange and pay for water and sewer service directly to {w.get('provider','the service provider')}."))
    else:
        sub.append(("Water and Sewer", "Landlord will provide water and sewer service; Tenant will reimburse Landlord as Variable Charges."))

    # Trash
    t = util.get("trash", {})
    if t.get("included_in_rent"):
        line = f"Landlord will provide regular trash removal service ({t.get('provider','')}); cost is included in Base Rent."
        if t.get("pickup_day"):
            line += f" Pickup day: {t['pickup_day']}."
        sub.append(("Trash Removal", line.strip()))
    else:
        sub.append(("Trash Removal", f"Tenant will arrange and pay for trash removal service directly to {t.get('provider','the service provider')}."))

    # Snow Removal
    s = util.get("snow_removal", {})
    if s.get("provided_by") == "landlord":
        sub.append(("Snow Removal", "Landlord will provide snow removal services."))
    elif s.get("provided_by") == "tenant":
        sub.append(("Snow Removal", "Snow removal on walkways, driveways, and other areas as required by local ordinance is the responsibility of Tenant."))
    else:
        sub.append(("Snow Removal", "Snow removal services are not provided."))

    # Landscaping
    l = util.get("landscaping", {})
    if l.get("provided_by") == "landlord":
        sub.append(("Landscaping", "Landlord will provide landscaping services."))
    elif l.get("provided_by") == "tenant":
        if l.get("tenant_addendum_required"):
            sub.append(("Landscaping", "Tenant is responsible for landscaping in accordance with the Tenant Yard Care Addendum."))
        else:
            sub.append(("Landscaping", "Tenant is responsible for landscaping."))
    else:
        sub.append(("Landscaping", "Landscaping services are not provided by Landlord."))

    # Telephone, Cable, Internet
    sub.append(("Telephone", label(util.get("telephone", {}), "telephone services")))
    sub.append(("Cable Television", label(util.get("cable_tv", {}), "cable or other premium television services")))
    sub.append(("Internet", label(util.get("internet", {}), "internet service")))

    for i, (name, line) in enumerate(sub, start=1):
        add_para(doc, f"1.6.{i}. {name}", bold=True)
        add_para(doc, line)


def write_1_7_security_deposit(doc, ctx):
    add_heading(doc, "1.7. Security Deposit", level=2)
    sd = ctx["tenancy"]["security_deposit"]
    bank = ctx["landlord"]["security_deposit_bank"]
    schedule = sd.get("payment_schedule", "lump_sum")
    if sd.get("already_held"):
        add_para(
            doc,
            f"Tenant's existing Security Deposit of {fmt_money(sd['amount'])} "
            f"(Security Deposit) shall continue to be held by Landlord under this "
            f"Lease.",
        )
    elif schedule == "installments":
        inst = sd.get("installments", {})
        first_due_at = inst.get("first_due", "start")
        when_label = "the Start Date" if first_due_at == "start" else "the signing of this Lease"
        add_para(
            doc,
            f"Tenant is required to pay a security deposit in the amount of "
            f"{fmt_money(sd['amount'])} (Security Deposit), payable in "
            f"{inst.get('count', 0)} consecutive equal monthly installments of "
            f"{fmt_money(inst.get('amount', 0))} beginning on {when_label}.",
        )
    else:
        add_para(
            doc,
            f"Tenant is required to pay a security deposit to Landlord when the "
            f"Lease is signed. The Security Deposit is {fmt_money(sd['amount'])} "
            f"(Security Deposit).",
        )
    add_para(
        doc,
        f"Landlord will hold the Security Deposit in an account at {bank['name']}, "
        f"located at {bank['address']}. Interest on the Security Deposit will be "
        f"retained by Landlord, unless required to be paid to Tenant under "
        f"applicable law. Section 2.4 of this Lease contains additional terms "
        f"relating to the Security Deposit.",
    )


def write_1_8_tenant_insurance(doc, ctx):
    add_heading(doc, "1.8. Tenant Insurance", level=2)
    ri = ctx["tenancy"].get("rental_insurance", {})
    if ri.get("required"):
        cov = ri.get("minimum_coverage")
        line = "Tenant is required to maintain renter's insurance during the Term"
        if cov:
            line += f" with minimum personal liability coverage of {fmt_money(cov)}."
        else:
            line += "."
        line += " Upon Landlord's request, Tenant will provide a certificate of insurance as evidence of the policy."
        add_para(doc, line)
    else:
        add_para(doc, "Tenant is not required to maintain renter's insurance during the Term.")


def write_1_9_parking(doc, ctx):
    add_heading(doc, "1.9. Parking", level=2)
    park = ctx["property"].get("parking", {})
    if not park.get("available") or not ctx["addenda"]["parking"]["include"]:
        add_para(doc, "No parking is provided as part of this Lease.")
        return
    if park.get("included_in_rent"):
        cost = "the cost of parking is included in Base Rent."
    else:
        cost = f"Parking Rent is {fmt_money(park.get('monthly_fee', 0))} per month, payable as Additional Rent."
    add_para(
        doc,
        f"Tenant may park in areas designated by Landlord (Parking Area), and "
        f"{cost} An addendum (Parking Addendum) is attached to this Lease setting "
        f"forth the specific terms of, and limitations on, Tenant's parking rights.",
    )


def write_1_10_storage(doc, ctx):
    add_heading(doc, "1.10. Storage Space", level=2)
    s = ctx["property"].get("storage", {})
    if not s.get("available"):
        add_para(doc, "No storage space is provided as part of this Lease.")
        return
    loc = s.get("location", "")
    if s.get("included_in_rent"):
        cost = "The rent for the Storage Space is included in Base Rent."
    else:
        cost = f"Tenant will pay {fmt_money(s.get('monthly_fee', 0))} per month for the Storage Space as Additional Rent."
    add_para(
        doc,
        f"Tenant may use the storage area exterior to the Property (Storage Space) "
        f"located at: {loc}. {cost} Tenant may not store any items which (i) pose a "
        f"threat of injury to any person or property, or (ii) would be determined "
        f"to be a hazardous substance or hazardous waste under any applicable law.",
    )


def write_1_11_pets(doc, ctx):
    add_heading(doc, "1.11. Pets", level=2)
    pets = ctx["addenda"].get("pets", {})
    if not pets.get("include"):
        add_para(doc, "Pets are not permitted on the Property.")
        return
    add_para(
        doc,
        "Tenant is only permitted to keep pet(s) on the Property that are "
        "identified in the addendum (Pet Addendum) attached to this Lease, and will "
        "comply with all terms of the Pet Addendum.",
    )
    if pets.get("insurance_required"):
        add_para(
            doc,
            "Tenant is required to carry renter's insurance which includes coverage "
            "for pet ownership and to provide proof of coverage upon Landlord's "
            "request.",
        )


def write_1_12_additional_rules(doc, ctx):
    add_heading(doc, "1.12. Additional Rules", level=2)
    if ctx["addenda"].get("rules", {}).get("include"):
        add_para(
            doc,
            "Tenant's use and occupancy of the Property is subject to the rules and "
            "regulations set forth in the Rules Addendum attached to this Lease.",
        )
    else:
        add_para(
            doc,
            "Tenant's use and occupancy of the Property is subject to all applicable "
            "law and the terms of this Lease.",
        )


def write_1_13_smoking(doc, ctx):
    add_heading(doc, "1.13. Smoking Policy", level=2)
    fee = ctx["landlord"]["defaults"].get("smoking_violation_fee", 250)
    add_para(
        doc,
        "Smoking means: (i) inhaling, exhaling, breathing, carrying, or possessing "
        "any lighted cigar, cigarette, pipe, or other lighted or heated tobacco or "
        "plant product intended for inhalation, including hookahs, whether natural "
        "or synthetic, in any manner or form; or (ii) use of an electronic smoking "
        "device which creates an aerosol or vapor, in any manner or form.",
    )
    add_para(doc, "Smoking is not allowed on the Property or anywhere on the Property grounds.")
    add_para(
        doc,
        f"In addition to any other remedies Landlord may have under this Lease or "
        f"at law, Landlord may charge Tenant a fee of up to {fmt_money(fee)} for a "
        f"second or subsequent violation of this smoking policy.",
    )


# ---------------------------------------------------------------------------
# Section 2: Additional Terms — mostly stable boilerplate
# ---------------------------------------------------------------------------

def write_section_2_additional_terms(doc, ctx):
    add_heading(doc, "2. Additional Terms", level=1, page_break_before=True)

    add_heading(doc, "2.1. Property Condition", level=2)
    add_para(
        doc,
        "Tenant has examined the Property, either in person or virtually, prior to "
        "signing this Lease and, as of the date of this Lease, is satisfied with its "
        "condition and appearance (Existing Condition). Landlord will deliver "
        "possession of the Property to Tenant on the Start Date in the same or "
        "better condition as the Existing Condition, except for ordinary wear and "
        "tear.",
    )

    add_heading(doc, "2.2. Possession", level=2)
    add_para(
        doc,
        "If Landlord cannot deliver possession of the Property to Tenant by the "
        "Start Date through no fault of Landlord, this Lease will continue in full "
        "force, but Tenant, as Tenant's sole remedy, will not be obligated to pay "
        "Monthly Rent (prorated based on a 30-day month) for the period that Tenant "
        "is unable to take possession. If Landlord fails to deliver possession by "
        "the 30th day following the Start Date, either party may terminate this "
        "Lease by written notice, in which case all amounts paid to Landlord by "
        "Tenant will be returned and both parties will be released from all "
        "obligations under this Lease.",
    )

    add_heading(doc, "2.3. Rent Payment", level=2)
    grace = ctx["tenancy"]["rent"].get("grace_period_days", 5)
    add_para(doc, "2.3.1. Payment Timing", bold=True)
    add_para(
        doc,
        "Tenant will pay Monthly Rent to Landlord, in advance, on the Monthly Rent "
        "Due Date of each month during the Term according to the payment details "
        "specified in Section 1.5. Variable Charges, if any, are payable according "
        "to the timeframe specified in Section 1.5.",
    )
    add_para(doc, "2.3.2. Late Payment", bold=True)
    add_para(
        doc,
        f"If Tenant fails to pay any Rent in full by the end of the day "
        f"{grace} day(s) after it is due, the Late Fee specified in Section 1.5 "
        f"will be assessed. Acceptance of late payment does not waive Landlord's "
        f"right to require payment of Monthly Rent in full on the date it is due.",
    )
    add_para(doc, "2.3.3. Returned Checks / Dishonored Payments", bold=True)
    add_para(
        doc,
        "If any payment under this Lease is returned for insufficient funds or "
        "otherwise fails, Landlord may require the dishonored payment be replaced "
        "by a cashier's check, certified check, or money order, and may require "
        "Tenant to pay any fee associated with the failed payment. If more than "
        "two of Tenant's payments during the Term are returned for insufficient "
        "funds, Landlord may require all future payments be made by cashier's "
        "check, certified check, or money order.",
    )

    add_heading(doc, "2.4. Security Deposit", level=2)
    add_para(doc, "2.4.1. Use of Security Deposit", bold=True)
    add_para(
        doc,
        "Subject to applicable law, Landlord may use the Security Deposit: "
        "(i) to remedy Tenant Defaults under this Lease, including past due Rent; "
        "(ii) to pay for costs incurred by Landlord to repair damages to the "
        "Property caused by Tenant, any Occupant, or any guest, beyond ordinary "
        "wear and tear; (iii) to pay cleaning costs incurred by Landlord to return "
        "the Property to the same level of cleanliness it was in at the Start "
        "Date; and (iv) for any other purpose permitted by applicable law "
        "(collectively, Deposit Claims). The Security Deposit will not relieve "
        "Tenant of any obligation to pay Rent due under this Lease prior to "
        "termination. If a Pet Damage Deposit is required, it will be considered "
        "a Security Deposit and subject to the terms of this Section.",
    )
    add_para(doc, "2.4.2. Return of Security Deposit", bold=True)
    add_para(
        doc,
        "The Security Deposit and any accrued interest (if required by applicable "
        "law), less any Deposit Claims (collectively, Security Deposit Balance), "
        "will be returned to Tenant within 30 days (or shorter period if required "
        "by applicable law) after Tenant vacates the Property upon expiration or "
        "earlier termination of this Lease. Any Deposit Claims will be described "
        "in an itemized statement provided with the Security Deposit Balance. "
        "Tenant will provide a forwarding address to Landlord for delivery of "
        "the Security Deposit Balance and itemized statement.",
    )

    add_heading(doc, "2.5. Tenant's Obligations", level=2)
    obligations = [
        ("2.5.1. Residential Use Only",
         "Tenant will use and occupy the Property for residential purposes only and "
         "will not use the Property for any non-residential, illegal, or otherwise "
         "inappropriate purpose, including any commercial purpose."),
        ("2.5.2. Permitted Occupants",
         "Subject to applicable law, the Property will not be occupied by anyone "
         "other than Tenants and Occupants identified in the Basic Terms."),
        ("2.5.3. No Disturbance or Nuisance",
         "Tenant will not, and will not permit any Occupant or guest to: (i) make "
         "any unreasonably loud or otherwise unreasonable use of the Property; "
         "(ii) allow any condition that poses a threat of injury to persons or "
         "property; or (iii) interfere with the rights, comfort, safety, or "
         "enjoyment of neighboring properties."),
        ("2.5.4. Utilities",
         "Tenant will not cause any utility to be interrupted during the Term and "
         "will provide Landlord with reasonable evidence that utilities for which "
         "Tenant is responsible have been paid upon Landlord's reasonable request."),
        ("2.5.5. Maintenance",
         "Tenant will: (i) keep and maintain the Property in a clean, safe, and "
         "sanitary condition; (ii) regularly dispose of garbage and waste in a "
         "clean and safe manner; (iii) use all appliances, fixtures, and equipment "
         "in a safe and reasonable manner; (iv) not obstruct doors and windows; "
         "and (v) maintain the Property in the same condition as it was delivered, "
         "except for ordinary wear and tear."),
        ("2.5.6. No Transfer",
         "Tenant will not sublease or assign all or any portion of the Property or "
         "this Lease without the prior written consent of Landlord. Tenant will "
         "not rent the Property, or any portion of the Property, through any "
         "short-term rental program (e.g., Airbnb, VRBO), and entry into any such "
         "agreement is cause for termination of this Lease."),
        ("2.5.7. No Alterations",
         "Tenant will not perform any alterations or improvements to the Property "
         "without the prior written consent of Landlord, in Landlord's sole "
         "discretion. Alterations include adding, changing, or removing "
         "appliances, fixtures, shelving, wallpaper, or wall paint, and the "
         "installation of new wiring, cabling, or equipment (except as required "
         "by applicable law)."),
        ("2.5.8. Joint Liability",
         "All individuals executing this Lease as Tenants will be jointly and "
         "severally liable for the performance of all obligations of a Tenant "
         "under this Lease."),
    ]
    for header, body in obligations:
        add_para(doc, header, bold=True)
        add_para(doc, body)

    add_heading(doc, "2.6. Landlord's Obligations", level=2)
    add_para(doc, "2.6.1. Services and Utilities", bold=True)
    add_para(
        doc,
        "Landlord will provide the services and utilities specified in Section 1.6 "
        "and as otherwise required under applicable law. Except as otherwise "
        "provided by applicable law, Tenant waives all liability of Landlord for "
        "any interruption or insufficiency of any service or utility resulting "
        "from causes beyond the reasonable control of Landlord.",
    )
    add_para(doc, "2.6.2. Maintenance and Repairs", bold=True)
    add_para(
        doc,
        "Subject to Tenant's duties under Section 2.5, Landlord will maintain the "
        "Property (including its structural elements, roof, and systems) in good "
        "order and repair. Landlord will be responsible for repairing the "
        "appliances, fixtures, and equipment located in the Property, except if "
        "any repairs are necessary as a result of improper use by Tenant, any "
        "Occupant, or a guest. Tenant will notify Landlord promptly in writing "
        "upon becoming aware of any condition requiring repair or maintenance. "
        "Landlord will undertake any required repairs reasonably promptly "
        "following receipt of notice.",
    )

    add_heading(doc, "2.7. Landlord's Access", level=2)
    add_para(
        doc,
        "Landlord, its agents, and contractors will have the right of reasonable "
        "access to the Property during normal business hours to perform "
        "maintenance and repair, and for the purpose of showing the Property to "
        "prospective tenants and purchasers. Tenant will be provided at least 24 "
        "hours' notice (or longer period if required by applicable law) prior to "
        "entry, except with Tenant's consent or in case of an emergency.",
    )

    add_heading(doc, "2.8. Surrender", level=2)
    add_para(
        doc,
        "Tenant will surrender possession of the Property and return all keys to "
        "Landlord upon the Expiration Date or earlier termination of this Lease. "
        "At surrender, the Property will be in the same condition as the Start "
        "Date, except for ordinary wear and tear, and otherwise in clean condition "
        "and free of all personal property of Tenant or any Occupant. Tenant's "
        "obligations under this Section will survive termination of this Lease.",
    )

    add_heading(doc, "2.9. Default", level=2)
    add_para(
        doc,
        "Tenant will be in default (Default) if Tenant: (i) fails to pay any Rent "
        "when due; or (ii) fails to comply with any other obligation or "
        "restriction in this Lease or any attached Addendum. If Tenant Defaults, "
        "Landlord may exercise all rights and remedies available under and in "
        "accordance with applicable law, including the right to: (i) terminate "
        "this Lease; (ii) regain possession of the Property through an eviction "
        "or similar process; (iii) recover from Tenant all unpaid Rent, including "
        "unpaid Monthly Rent, Additional Rent, Late Fees, and (if applicable) "
        "holdover Rent for the period prior to Tenant's delivery of possession of "
        "the Property to Landlord; (iv) recover all Rent payable under this Lease "
        "for the period from the date of termination through the stated Expiration "
        "Date (provided that Landlord has used reasonable efforts to find a "
        "replacement tenant), less the amount Landlord is able to collect from "
        "any replacement tenants for that period; and (v) recover all reasonable "
        "costs and expenses incurred by Landlord in repairing any damage to the "
        "Property caused by improper use by Tenant, any Occupant, or any guest of "
        "Tenant or an Occupant, less any amounts obtained from the Security "
        "Deposit. To the extent permitted under applicable law, Landlord may also "
        "recover Landlord's court costs and reasonable attorneys' fees and "
        "expenses incurred in connection with any legal proceedings against "
        "Tenant. To the extent required by applicable law, Landlord will use "
        "reasonable efforts to mitigate any damages resulting from Tenant Default.",
    )

    add_heading(doc, "2.10. Notices", level=2)
    add_para(
        doc,
        "Any notice of termination of this Lease, notice of Default by Tenant "
        "under this Lease, notice of eviction by Landlord, notice of disposition "
        "of Tenant's property, or any other notice required to be given in "
        "writing under applicable law (Material Notices) will be in writing and "
        "sent to Tenant and Landlord at the applicable address set forth in "
        "Section 2.16 or as otherwise required under applicable law. Except for "
        "Material Notices, all other written notices may be delivered to the "
        "other party at the email or physical address specified in Section 2.16, "
        "or by other electronic means agreed to by the parties. Either party may "
        "update its email or physical address by sending written notice to the "
        "other party.",
    )

    add_heading(doc, "2.11. Casualty Damage", level=2)
    add_para(
        doc,
        "If the Property is rendered uninhabitable by fire, storm, or other "
        "casualty, this Lease will terminate upon thirty (30) days' written "
        "notice. Tenant will pay Landlord only the Rent for the period prior to "
        "the casualty, and Landlord will return to Tenant any Rent paid for the "
        "period after the casualty unless Tenant remains in possession of the "
        "Property after the casualty. Landlord's written notice to vacate must "
        "indicate that (i) the Property has been certified or condemned as "
        "uninhabitable by a local agency charged with the authority to issue such "
        "an order; and (ii) continued habitation of the Property would subject "
        "Landlord to civil or criminal penalties. If the local agency's order does "
        "not permit Landlord to provide thirty (30) days' advance written notice, "
        "Landlord must provide as much advance written notice as is possible while "
        "complying with the order.",
    )
    add_para(
        doc,
        "If a portion of the Property remains habitable, this Lease will continue "
        "but Monthly Rent will be adjusted proportionally based on the proportion "
        "of the Property still habitable by Tenant, until the damaged portion has "
        "been restored. Nonessential elements (including decks and porches) of "
        "the Property will not be counted in determining the habitable portions "
        "of the Property.",
    )
    add_para(
        doc,
        "If the Property is damaged or destroyed by fire or other casualty "
        "resulting from any negligent act by Tenant, any Occupant, or any guest "
        "of Tenant or any Occupant, Tenant will be liable to Landlord for the "
        "costs of repairing the damage.",
    )

    add_heading(doc, "2.12. Tenant's Property", level=2)
    add_para(
        doc,
        "Tenant acknowledges that Landlord's insurance does not cover loss or "
        "damage to any of Tenant's personal property and that Landlord will not "
        "be liable for any such damage. Even if no policy of renter's insurance "
        "is required, Landlord recommends that Tenant obtain renter's insurance.",
    )

    add_heading(doc, "2.13. General", level=2)
    add_para(
        doc,
        "This Lease will be governed by the laws of the State of Washington and "
        "any additional laws of the city or county in which the Property is "
        "located. This Lease will be binding on and inure to the benefit of all "
        "permitted heirs, legal representatives, and assigns of the parties. This "
        "Lease, along with the attached Addenda and legal disclosures, contains "
        "the entire agreement between Landlord and Tenant and may not be changed "
        "except in writing signed by all parties. If any provision of this Lease "
        "is found invalid or unenforceable, all other provisions will remain "
        "binding to the maximum extent permitted by applicable law.",
    )

    add_heading(doc, "2.14. Disclosures / Addenda", level=2)
    add_para(
        doc,
        "Tenant acknowledges that the legal disclosures and addenda (Addenda) "
        "attached to this Lease are part of the legal agreement between the "
        "parties. Tenant will comply with all applicable rules and regulations set "
        "out in the attached Addenda. The terms of this Lease will control in the "
        "event of any conflict between the terms of any Addenda and the terms of "
        "the Lease.",
    )

    add_heading(doc, "2.15. Execution", level=2)
    add_para(
        doc,
        "All individuals indicated as Tenant will sign this Lease and related "
        "Addenda where indicated. Each of Landlord and Tenant consents to the "
        "other party's execution of this Lease by electronic signature. Delivery "
        "by facsimile, electronic means, or as a digital copy will have the same "
        "full force and effect as a manually executed original.",
    )


def write_contact_info(doc, ctx):
    add_heading(doc, "2.16. Contact Information", level=2)
    add_para(doc, "2.16.1. Tenant", bold=True)
    add_para(doc, "Tenant's address is required for notice prior to the Start Date. Notices after the Start Date will be made to the Property.")

    table = doc.add_table(rows=1 + len(ctx["tenancy"]["tenants"]), cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Tenant Name", "Address for Notice (Pre-Start)", "Phone", "Email"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for i, t in enumerate(ctx["tenancy"]["tenants"], start=1):
        cells = table.rows[i].cells
        cells[0].text = t.get("name", "")
        cells[1].text = t.get("address_for_notice_pre_start", "")
        cells[2].text = t.get("phone", "")
        cells[3].text = t.get("email", "")

    add_para(doc, "", size=4)
    add_para(doc, "2.16.2. Landlord", bold=True)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Landlord Name", "Address for Notice", "Phone", "Email"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
    cells = table.rows[1].cells
    cells[0].text = ctx["landlord"]["name"]
    cells[1].text = fmt_address(ctx["landlord"]["address"])
    cells[2].text = ctx["landlord"].get("phone", "")
    cells[3].text = ctx["landlord"].get("email", "")


def write_main_signature(doc, ctx):
    add_para(doc, "", size=10)
    add_para(
        doc,
        "IN WITNESS WHEREOF, Tenant and Landlord have executed this Lease as of "
        "the date of the last signature below.",
    )
    add_signature_block(doc, "Tenant", [t["name"] for t in ctx["tenancy"]["tenants"]])
    add_signature_block(doc, "Landlord", [ctx["landlord"]["name"]])


# ---------------------------------------------------------------------------
# Addenda
# ---------------------------------------------------------------------------

def write_pet_addendum(doc, ctx):
    add_heading(doc, "Pet Addendum", level=1, page_break_before=True)
    add_para(
        doc,
        "This Pet Addendum is attached to and made a part of the Lease between "
        "Landlord and Tenant for the Property dated as of the date hereof "
        "(Lease). All capitalized terms used in this Addendum have the meanings "
        "given such terms in the Lease.",
    )

    pets = ctx["addenda"]["pets"]

    add_heading(doc, "A. Permission to Have Pet(s) / Identification", level=2)
    add_para(
        doc,
        "Subject to Tenant's compliance with this Addendum, Tenant is granted "
        "permission to keep only the following pet(s) at the Property during the "
        "Term of the Lease (each, a Pet):",
    )
    table = doc.add_table(rows=1 + len(pets["pets"]), cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].paragraphs[0].add_run("Pet Name").bold = True
    hdr[1].paragraphs[0].add_run("Description").bold = True
    for i, pet in enumerate(pets["pets"], start=1):
        cells = table.rows[i].cells
        cells[0].text = pet["name"]
        cells[1].text = pet["description"]

    # Pet financial terms
    fin_lines = []
    if pets.get("pet_rent_monthly"):
        fin_lines.append(
            f"Pet Rent: {fmt_money(pets['pet_rent_monthly'])} per month, payable as Additional Rent."
        )
    if pets.get("non_refundable_pet_fee"):
        fin_lines.append(
            f"Non-Refundable Pet Fee: {fmt_money(pets['non_refundable_pet_fee'])}, due at signing of this Addendum."
        )
    if pets.get("refundable_pet_deposit"):
        fin_lines.append(
            f"Refundable Pet Deposit: {fmt_money(pets['refundable_pet_deposit'])}, "
            f"held as part of the Security Deposit."
        )
    if fin_lines:
        add_heading(doc, "B. Pet Financial Terms", level=2)
        for line in fin_lines:
            add_para(doc, "• " + line)

    add_heading(doc, "C. Rules, Regulations, and Agreements", level=2)
    rules = [
        "Tenant represents that each Pet is properly licensed and vaccinated as required by applicable law, and Tenant agrees to keep all licensure and vaccinations current. Tenant will provide proof of licensing and vaccination upon Landlord's request.",
        "Tenant represents that each Pet is housebroken, has no vicious tendencies, and has no history of threatening or causing harm to persons.",
        "Pets must be kept on a leash or otherwise under the direct physical control of Tenant or another person at all times when outside the Property.",
        "Tenant will clean up after each Pet and properly dispose of all waste.",
        "Tenant will maintain the Property free of odor and stain from any Pet, and will use prompt and diligent efforts to prevent and eliminate any infestation of pests (e.g., fleas) associated with any Pet.",
        "Tenant is responsible for controlling and minimizing noise caused by any Pet. Unreasonable levels of noise that interfere with the quiet enjoyment of neighboring tenants or properties is prohibited.",
        "Tenant is responsible for and will be charged for any damage to the Property caused by any Pet, including damage to floors, carpets, drapes, screens, landscaping, fencing, and any odors due to the presence of a Pet.",
        "Tenant will indemnify and hold Landlord harmless from all liability, claims, demands, damages, and costs for injuries to persons or property in connection with Tenant's Pet(s).",
        "If Landlord receives a complaint or otherwise has reasonable belief that the conduct or condition of a Pet constitutes a nuisance under state or local law or otherwise poses a threat to the safety or health of others, Landlord may inspect the Property and, if Landlord determines that the Pet constitutes a nuisance or threat, Landlord may revoke the permission granted under this Addendum and order Tenant to remove the subject Pet from the Property within 48 hours.",
    ]
    if pets.get("insurance_required"):
        rules.insert(
            1,
            "Tenant is required to maintain renter's insurance with coverage for pet ownership, and to provide proof of coverage upon Landlord's request.",
        )
    numbered_list(doc, rules)

    add_para(doc, "The violation of any provision of this Pet Addendum will constitute a Default under the Lease.")
    add_signature_block(doc, "Tenant", [t["name"] for t in ctx["tenancy"]["tenants"]])
    add_signature_block(doc, "Landlord", [ctx["landlord"]["name"]])


def write_parking_addendum(doc, ctx):
    add_heading(doc, "Parking Addendum", level=1, page_break_before=True)
    add_para(
        doc,
        "This Parking Addendum is attached to and made a part of the Lease between "
        "Landlord and Tenant for the Property dated as of the date hereof (Lease). "
        "All capitalized terms used in this Addendum have the meanings given such "
        "terms in the Lease.",
    )
    add_para(
        doc,
        "Tenant may park vehicles in the parking area on the grounds of the "
        "Property during the Lease Term. Landlord may require all vehicles parking "
        "in the Parking Area to be registered with Landlord (Registered Vehicles), "
        "including providing license plate number and owner's contact information. "
        "No vehicles other than Registered Vehicles may be parked in the parking "
        "area by Tenant, any Occupant, or any guest. If Tenant replaces any "
        "Registered Vehicle, Tenant must notify Landlord and provide updated "
        "vehicle information prior to parking the replacement vehicle.",
    )
    rules = [
        "All posted parking and traffic regulations must be obeyed.",
        "The parking area will be used only to park motor vehicles and for loading or unloading of motor vehicles.",
        "All ordinances regarding fire lanes must be obeyed. Any vehicle parked outside the parking area, in a fire lane, blocking a fire hydrant, refuse container, another vehicle, sidewalk, or lawn, or otherwise illegally parked may be towed at the vehicle owner's expense.",
        "Only operable passenger vehicles (including pick-up trucks) that reasonably fit in a designated parking space may use the parking area. Commercial vehicles, recreational vehicles, boats, trailers, and other oversized vehicles are not permitted.",
        "Landlord may remove any vehicle at the owner's expense if it appears abandoned or inoperable, lacks an inspection sticker or license plates, or has expired inspection or registration.",
        "Repairs to vehicles are prohibited in the parking area or on Property grounds, except for emergency repairs.",
        "Vehicles may be washed only in designated areas. If there is no designated area, washing vehicles is not allowed on the grounds of the Property.",
        "Tenant's use of the parking spaces and parking area is at Tenant's own risk. Landlord does not provide security for the parking area and is not liable for damage to or theft of any vehicle or property from any vehicle.",
    ]
    numbered_list(doc, rules)
    add_para(doc, "The violation of any provision of this Parking Addendum will constitute a Default under the Lease.")
    add_signature_block(doc, "Tenant", [t["name"] for t in ctx["tenancy"]["tenants"]])
    add_signature_block(doc, "Landlord", [ctx["landlord"]["name"]])


def write_rules_addendum(doc, ctx):
    add_heading(doc, "Rules Addendum", level=1, page_break_before=True)
    add_para(
        doc,
        "This Rules Addendum is attached to and made a part of the Lease between "
        "Landlord and Tenant for the Property dated as of the date hereof (Lease). "
        "All capitalized terms used in this Addendum have the meanings given such "
        "terms in the Lease.",
    )

    sections = [
        ("Actions of Residents", [
            "Tenant will dispose of trash only in designated areas.",
            "Tenant will promptly report any repair or maintenance problems to Landlord.",
        ]),
        ("Keys", [
            "At delivery of possession to the Property, Tenant will sign and deliver to Landlord a receipt identifying the locks associated with the keys provided to Tenant and the number of each type of key.",
            "When Tenant vacates the Property at the end of the Term, Tenant will return all keys provided to Tenant by Landlord. If Tenant fails to return all keys, Landlord may re-key the applicable locks at Tenant's expense.",
            "If Tenant loses a key or requires duplicates, Tenant must notify Landlord and bear the cost. Tenant may not duplicate keys without Landlord's consent.",
        ]),
        ("Use of Premises", [
            "Tenant may not fasten anything to the fixtures, appliances, or to the interior or exterior of the Property without Landlord's written consent.",
            "Any balcony or porch may not be altered by Tenant or used to store personal belongings.",
            "Tenant will comply with all weight restrictions on balconies and porches.",
            "No laundry or other items will be hung from any window, balcony, or porch.",
            "Tenant will not bring anything onto the Property which could increase the risk of fire (e.g., flammable chemicals).",
            "Tenant will not place any sign, advertisement, or notice visible from outside the Property.",
            "Tenant will not add or change any locks without Landlord's prior written consent.",
            "Waterbeds and other water furniture are prohibited. Unusually heavy items (e.g., pianos, safes) require Landlord's written consent.",
        ]),
    ]
    for title, items in sections:
        add_heading(doc, title, level=2)
        numbered_list(doc, items)

    custom = ctx["addenda"]["rules"].get("custom_rules", [])
    if custom:
        add_heading(doc, "Additional Rules", level=2)
        numbered_list(doc, custom)

    add_para(doc, "The violation of any provision of this Rules Addendum will constitute a Default under the Lease.")
    add_signature_block(doc, "Tenant", [t["name"] for t in ctx["tenancy"]["tenants"]])
    add_signature_block(doc, "Landlord", [ctx["landlord"]["name"]])


def write_lawn_care_addendum(doc, ctx):
    add_heading(doc, "Tenant Yard Care Addendum", level=1, page_break_before=True)
    add_para(
        doc,
        f"This Addendum is attached to and made a part of the Lease between "
        f"{ctx['landlord']['name']} (Landlord) and {ctx['tenant_names_joined']} "
        f"(Tenant) for the property located at {ctx['property_address_oneline']} "
        f"(Property), dated as of the date hereof. All capitalized terms used in "
        f"this Addendum have the meanings given such terms in the Lease.",
    )

    add_heading(doc, "Purpose", level=2)
    add_para(
        doc,
        "This Addendum sets forth the responsibilities of Tenant and Landlord "
        "regarding the care and maintenance of the yard and landscaping at the "
        "Property.",
    )

    add_heading(doc, "Tenant Responsibilities", level=2)
    add_para(doc, "Tenant agrees to perform the following yard care tasks:")
    numbered_list(doc, [
        "Mowing: Tenant shall mow the lawn regularly, maintaining a grass height of no more than five inches.",
        "Watering: Tenant shall water the lawn, plants, and shrubs as needed to maintain their health and appearance.",
        "Weeding: Tenant shall keep the lawn, flower beds, and other landscaped areas free of weeds.",
        "Trimming: Tenant shall trim hedges, shrubs, and trees as needed to maintain their shape and size.",
        "Raking: Tenant shall rake leaves and other debris from the lawn and landscaped areas as needed.",
        "Snow Removal: Tenant shall remove snow and ice from sidewalks, driveways, and other areas as required by local ordinances.",
    ])

    add_heading(doc, "Landlord Responsibilities", level=2)
    numbered_list(doc, [
        "Fertilizing: Landlord shall fertilize the lawn and landscaped areas as needed.",
        "Pest Control: Landlord shall provide pest control services for the lawn and landscaped areas as needed.",
        "Major Landscaping: Landlord shall be responsible for major landscaping projects, such as tree removal or installation of new plants.",
    ])

    add_heading(doc, "Additional Provisions", level=2)
    equipment_provider = ctx["addenda"]["lawn_care"].get("equipment_provider", "landlord")
    if equipment_provider == "tenant":
        equipment_line = (
            "Equipment: Tenant shall provide their own lawn care equipment, unless "
            "otherwise agreed upon in writing."
        )
    else:
        equipment_line = (
            "Equipment: Landlord shall provide lawn care equipment, unless "
            "otherwise agreed upon in writing."
        )
    numbered_list(doc, [
        equipment_line,
        "Damage: Tenant shall be responsible for any damage to the lawn or landscaping caused by their negligence or misuse.",
        "Inspections: Landlord or Landlord's agent shall have the right to inspect the yard and landscaping at reasonable times with reasonable notice.",
    ])

    add_heading(doc, "Consequences of Non-Compliance", level=2)
    add_para(
        doc,
        "Failure by Tenant to fulfill the yard care responsibilities outlined in "
        "this Addendum may result in:",
    )
    numbered_list(doc, [
        "Written Notice: Landlord shall provide Tenant with written notice of the non-compliance and a 10-day period to remedy the issue.",
        "Landlord Remedy: If Tenant fails to remedy the issue within the 10-day period, Landlord may, at Landlord's option, perform the necessary yard care tasks and charge Tenant for the cost as Additional Rent.",
        "Default: Repeated or significant non-compliance with this Addendum will constitute a Default under the Lease.",
    ])

    add_para(doc, "The violation of any provision of this Tenant Yard Care Addendum will constitute a Default under the Lease.")
    add_signature_block(doc, "Tenant", [t["name"] for t in ctx["tenancy"]["tenants"]])
    add_signature_block(doc, "Landlord", [ctx["landlord"]["name"]])


def write_lead_paint_disclosure(doc, ctx):
    add_heading(doc, "Lead-Based Paint Hazard Disclosure", level=1, page_break_before=True)
    add_heading(doc, "Lead Warning Statement", level=2)
    add_para(
        doc,
        "Housing built before 1978 may contain lead-based paint. Lead from paint, "
        "paint chips, and dust can pose health hazards if not managed properly. "
        "Lead exposure is especially harmful to young children and pregnant "
        "women. Before renting pre-1978 housing, lessors must disclose the "
        "presence of known lead-based paint and/or lead-based paint hazards in "
        "the dwelling. Lessees must also receive a federally approved pamphlet "
        "on lead poisoning prevention.",
    )
    add_heading(doc, "Landlord's Disclosures", level=2)
    add_para(doc, "Landlord has no knowledge of lead-based paint and/or lead-based paint hazards in the housing.")
    add_para(doc, "Landlord has no records or reports pertaining to lead-based paint and/or lead-based paint hazards in the housing.")
    add_heading(doc, "Tenant's Acknowledgements", level=2)
    add_para(
        doc,
        "By signing below, Tenant acknowledges receipt of copies of all "
        "information listed above and acknowledges receipt of the pamphlet "
        "Protect Your Family from Lead in Your Home, a copy of which is attached "
        "to this Lease.",
    )
    add_signature_block(doc, "Tenant", [t["name"] for t in ctx["tenancy"]["tenants"]])
    add_signature_block(doc, "Landlord", [ctx["landlord"]["name"]])


def write_mold_addendum(doc, ctx):
    add_heading(doc, "Mold Addendum", level=1, page_break_before=True)
    add_heading(doc, "Tenant's Acknowledgement", level=2)
    add_para(
        doc,
        "By signing below, Tenant acknowledges receipt of the pamphlet entitled "
        "\"A Brief Guide to Mold, Moisture, and Your Home,\" a copy of which is "
        "attached to this Lease.",
    )
    add_signature_block(doc, "Tenant", [t["name"] for t in ctx["tenancy"]["tenants"]])
    add_signature_block(doc, "Landlord", [ctx["landlord"]["name"]])


def write_fire_safety_addendum(doc, ctx):
    add_heading(doc, "Fire Safety Addendum", level=1, page_break_before=True)
    add_para(
        doc,
        "The Washington Residential Landlord-Tenant Act requires a landlord to "
        "notify all tenants that the property is equipped with a smoke detection "
        "device, of Tenant's responsibility to maintain the smoke detection "
        "device in proper operating condition, and of penalties for Tenant's "
        "failure to comply with their maintenance responsibility.",
    )
    sd = ctx["property"].get("smoke_detectors", {})
    if sd.get("installed"):
        battery = sd.get("battery_operated", True)
        line = (
            "Landlord has installed one or more operational smoke detectors in the "
            "Property. Tenant will regularly (not less often than once per month) "
            "test the smoke detectors"
        )
        if battery:
            line += (
                ", and, because the smoke detectors are battery operated, Tenant will "
                "replace the batteries every 6 months, or more often as needed, and "
                "otherwise maintain the smoke detectors as specified by the "
                "manufacturer."
            )
        else:
            line += (
                ", and otherwise maintain the smoke detectors as specified by the "
                "manufacturer."
            )
        line += (
            " Tenant will not disable any smoke detector. Tenant must report to "
            "Landlord any failure of a smoke detector to function properly. Failure "
            "to maintain the smoke detectors is punishable by a fine of up to $200 "
            "and is also a violation of this Lease."
        )
        add_para(doc, line)
    if not ctx["property"].get("fire_alarm_system"):
        add_para(doc, "The Property does not have a fire alarm system.")
    add_signature_block(doc, "Tenant", [t["name"] for t in ctx["tenancy"]["tenants"]])
    add_signature_block(doc, "Landlord", [ctx["landlord"]["name"]])


def ordinal(n):
    if 11 <= (n % 100) <= 13:
        return f"{n}th"
    return f"{n}{ {1:'st',2:'nd',3:'rd'}.get(n%10,'th') }"


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) != 5:
        print("Usage: generate_lease.py <landlord.json> <property.json> <tenancy.json> <output.docx>")
        sys.exit(2)
    landlord = json.loads(Path(sys.argv[1]).read_text())
    property_ = json.loads(Path(sys.argv[2]).read_text())
    tenancy = json.loads(Path(sys.argv[3]).read_text())
    generate(landlord, property_, tenancy, sys.argv[4])
    print(f"Wrote {sys.argv[4]}")


if __name__ == "__main__":
    main()
