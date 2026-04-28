#!/usr/bin/env python3
"""Generate a property condition report (move-in / move-out walkthrough form).

Usage:
    python3 generate_condition_report.py <landlord.json> <property.json> <tenancy.json> <output.docx>

Produces a clean, fresh form (not derived from any copyrighted template). The room
list is driven by the property profile's `rooms_for_condition_report` field.
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from lease_helpers import (
    fmt_address, join_names,
    add_heading, add_para, set_cell_shading,
)


# Items checked in every room. The wizard could later allow per-room overrides.
DEFAULT_ROOM_ITEMS = [
    "Ceiling",
    "Walls",
    "Floor / Flooring",
    "Doors / Locks",
    "Windows / Tracks / Screens",
    "Window Coverings",
    "Light Fixtures",
    "Outlets / Switches",
    "Other:",
]

# Some rooms have specialized items.
ROOM_ITEM_OVERRIDES = {
    "Kitchen": [
        "Cabinets & Counters",
        "Sink / Faucet",
        "Refrigerator (make/model/serial)",
        "Stove / Oven (make/model/serial)",
        "Microwave",
        "Dishwasher (make/model/serial)",
        "Disposal",
        "Hood / Filter / Fan",
        "Walls",
        "Ceiling",
        "Floor / Flooring",
        "Doors / Locks",
        "Windows / Tracks / Screens",
        "Light Fixtures",
        "Other:",
    ],
    "Bathroom 1": [
        "Cabinets / Counters",
        "Sink / Faucet",
        "Toilet",
        "Tub / Shower",
        "Exhaust Fan / Heater",
        "Mirror",
        "Walls",
        "Ceiling",
        "Floor / Flooring",
        "Doors / Locks",
        "Other:",
    ],
    "Bathroom 2": [
        "Cabinets / Counters",
        "Sink / Faucet",
        "Toilet",
        "Tub / Shower",
        "Exhaust Fan / Heater",
        "Mirror",
        "Walls",
        "Ceiling",
        "Floor / Flooring",
        "Doors / Locks",
        "Other:",
    ],
    "Grounds": [
        "Fences / Gates",
        "Landscape",
        "Lawn",
        "Walkways / Driveway",
        "Mailbox",
        "Other:",
    ],
    "Exterior": [
        "Siding / Paint",
        "Roof (visible condition)",
        "Gutters / Downspouts",
        "Foundation",
        "Decks / Porches",
        "Other:",
    ],
    "Basement": [
        "Walls",
        "Floor",
        "Ceiling",
        "Furnace (make/model/filter location)",
        "Water Heater",
        "Washer (make/model/serial)",
        "Dryer (make/model/serial)",
        "Sump Pump",
        "Other:",
    ],
    "Garage / Storage Shed": [
        "Doors / Openers",
        "Walls",
        "Floor",
        "Ceiling",
        "Light Fixtures",
        "Outlets",
        "Other:",
    ],
}


def items_for_room(room):
    return ROOM_ITEM_OVERRIDES.get(room, DEFAULT_ROOM_ITEMS)


def generate(landlord, property_, tenancy, output_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Cover / header
    add_heading(doc, "PROPERTY CONDITION REPORT", level=0, center=True)
    add_para(doc, "Move-in and move-out walkthrough", size=11, align="center", italic=True)
    add_para(doc, "")
    add_para(
        doc,
        "This Property Condition Report is part of the Lease between Landlord and "
        "Tenant. Tenant agrees to inspect the Property within seven (7) days of the "
        "Start Date and to mark the move-in condition of each item listed below. "
        "Landlord and Tenant will sign and date this report after move-in inspection "
        "is complete and again at move-out. Anything not noted at move-in is "
        "presumed to be in good condition.",
    )
    add_para(
        doc,
        "Rate each item as G (Good), F (Fair), or P (Poor). Use the Comments column "
        "to describe any specific condition; attach photographs for clarity where "
        "useful.",
    )

    add_heading(doc, "Unit Information", level=1)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.autofit = False
    rows = [
        ("Property Address", fmt_address(property_["address"])),
        ("Tenant(s)", join_names([t["name"] for t in tenancy["tenants"]])),
        ("Move-in Date", tenancy["term"].get("start_date", "")),
        ("Move-out Date", ""),
    ]
    for i, (k, v) in enumerate(rows):
        cells = info_table.rows[i].cells
        cells[0].width = Inches(2.0)
        cells[1].width = Inches(4.5)
        cells[0].paragraphs[0].add_run(k).bold = True
        cells[1].paragraphs[0].add_run(str(v))

    # Per-room walkthrough tables
    rooms = property_.get("rooms_for_condition_report", [])
    if not rooms:
        rooms = ["Living Room", "Kitchen", "Bathroom 1", "Bedroom 1", "Exterior"]

    for room in rooms:
        doc.add_paragraph()  # spacing
        add_heading(doc, room, level=2)
        items = items_for_room(room)
        # Table: Item | Move-in (G/F/P) | Move-in Comments | Move-out (G/F/P) | Move-out Comments
        table = doc.add_table(rows=1 + len(items), cols=5)
        table.style = "Light Grid Accent 1"
        table.autofit = False

        col_widths = [Inches(1.6), Inches(0.6), Inches(2.3), Inches(0.6), Inches(2.3)]
        headers = ["Item", "Move-in G/F/P", "Move-in Comments", "Move-out G/F/P", "Move-out Comments"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.width = col_widths[i]
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, "E7EDF5")
        for ri, item in enumerate(items, start=1):
            row = table.rows[ri]
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = w
            row.cells[0].paragraphs[0].add_run(item).font.size = Pt(10)

    # Pages get long; add a clear summary section
    doc.add_paragraph()
    add_heading(doc, "Summary of Existing Damage at Move-in", level=1)
    add_para(
        doc,
        "Use the space below to summarize any pre-existing damage observed at "
        "move-in that is not adequately captured in the room-by-room tables above. "
        "Attach additional pages or photographs as needed.",
    )
    for _ in range(8):
        add_para(doc, "_" * 90, size=10)

    add_heading(doc, "Summary of Damage at Move-out", level=1)
    add_para(
        doc,
        "To be completed by Landlord (with Tenant's input where possible) at move-out.",
    )
    for _ in range(8):
        add_para(doc, "_" * 90, size=10)

    # Acknowledgement / signatures
    doc.add_page_break()
    add_heading(doc, "Acknowledgement", level=1)
    add_para(
        doc,
        "Tenant and Landlord acknowledge that they have inspected the Property and "
        "that this report accurately reflects the condition of the Property at "
        "move-in. Tenant understands that conditions not noted on this report at "
        "move-in are presumed to be in good condition and may be considered Tenant "
        "damage at move-out, beyond ordinary wear and tear.",
    )

    add_para(doc, "")
    add_para(doc, "MOVE-IN", bold=True, size=11)
    for name in [t["name"] for t in tenancy["tenants"]]:
        add_para(doc, "")
        add_para(doc, "_" * 50, size=10)
        add_para(doc, f"Tenant: {name}    Date: _______________", size=10)
    add_para(doc, "")
    add_para(doc, "_" * 50, size=10)
    add_para(doc, f"Landlord: {landlord['name']}    Date: _______________", size=10)

    add_para(doc, "")
    add_para(doc, "MOVE-OUT", bold=True, size=11)
    for name in [t["name"] for t in tenancy["tenants"]]:
        add_para(doc, "")
        add_para(doc, "_" * 50, size=10)
        add_para(doc, f"Tenant: {name}    Date: _______________", size=10)
    add_para(doc, "")
    add_para(doc, "_" * 50, size=10)
    add_para(doc, f"Landlord: {landlord['name']}    Date: _______________", size=10)

    doc.save(output_path)


def main():
    if len(sys.argv) != 5:
        print("Usage: generate_condition_report.py <landlord.json> <property.json> <tenancy.json> <output.docx>")
        sys.exit(2)
    landlord = json.loads(Path(sys.argv[1]).read_text())
    property_ = json.loads(Path(sys.argv[2]).read_text())
    tenancy = json.loads(Path(sys.argv[3]).read_text())
    generate(landlord, property_, tenancy, sys.argv[4])
    print(f"Wrote {sys.argv[4]}")


if __name__ == "__main__":
    main()
