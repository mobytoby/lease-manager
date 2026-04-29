"""Shared helpers for lease document generation."""

from datetime import date, datetime
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def fmt_date(s):
    if not s:
        return ""
    if isinstance(s, (date, datetime)):
        d = s
    else:
        d = datetime.strptime(s, "%Y-%m-%d").date()
    return d.strftime("%B %-d, %Y")


def fmt_money(n):
    if n is None:
        return ""
    try:
        return "${:,.2f}".format(float(n))
    except (TypeError, ValueError):
        return str(n)


def fmt_address(addr):
    if not addr:
        return ""
    parts = [addr.get("street")]
    if addr.get("unit"):
        parts.append(addr["unit"])
    cs = ", ".join(p for p in [addr.get("city"), addr.get("state")] if p)
    if addr.get("zip"):
        cs = cs + " " + addr["zip"]
    parts.append(cs)
    return ", ".join(p for p in parts if p)


def join_names(names, conj="and"):
    if not names:
        return ""
    if len(names) == 1:
        return names[0]
    if len(names) == 2:
        return f"{names[0]} {conj} {names[1]}"
    return ", ".join(names[:-1]) + f", {conj} {names[-1]}"


def add_heading(doc, text, level=1, center=False, page_break_before=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(11)
    else:
        run.font.size = Pt(10)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if page_break_before:
        # Attach the page break to the heading paragraph rather than using a
        # standalone add_page_break() call.  A paragraph-attached break cannot
        # produce a blank page the way a free-standing break paragraph can.
        pPr = p._p.get_or_add_pPr()
        pb = OxmlElement("w:pageBreakBefore")
        pPr.append(pb)
        # Also keep the heading glued to the first line of content below it.
        p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text, *, bold=False, italic=False, size=10, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def _docusign_anchor(para, tag):
    """Append an invisible DocuSign anchor tag to an existing paragraph.

    The tag is rendered as 1pt white text so it's invisible to readers but
    detectable by DocuSign's Auto Place parser when the PDF is uploaded.
    Standard tags: \\s{n} signature, \\i{n} initials, \\d{n} date signed.
    """
    run = para.add_run(tag)
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Suppress spell-check so Word doesn't draw a squiggle over the tag
    # (squiggles render independently of font color and would be visible)
    rPr = run._r.get_or_add_rPr()
    noProof = OxmlElement("w:noProof")
    rPr.append(noProof)


def add_signature_block(doc, label, names, signer_start=None):
    """Add a signature block that always renders on a single page.

    signer_start: 1-based index of the first signer in this block. When set,
                  invisible DocuSign anchor tags are embedded so Auto Place
                  positions signature and date fields on upload.
                  Tenants are signers 1..N; landlord is N+1.
    """
    paras = []
    paras.append(add_para(doc, "", size=10))
    paras.append(add_para(doc, f"{label.upper()} SIGNATURE", bold=True, size=11))
    for i, name in enumerate(names):
        signer_num = (signer_start + i) if signer_start is not None else None
        paras.append(add_para(doc, "", size=10))

        sig_line = add_para(doc, "_" * 50, size=10)
        if signer_num is not None:
            _docusign_anchor(sig_line, f"\\s{signer_num}\\")
        paras.append(sig_line)

        paras.append(add_para(doc, name, size=10))

        date_line = add_para(doc, "Date: ____________________", size=10)
        if signer_num is not None:
            _docusign_anchor(date_line, f"\\d{signer_num}\\")
        paras.append(date_line)

    # Keep the whole block on one page: set keep_with_next on every line
    # except the last so Word won't split the block across a page boundary.
    for p in paras[:-1]:
        p.paragraph_format.keep_with_next = True


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)
